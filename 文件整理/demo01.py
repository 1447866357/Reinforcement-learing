# 1.找文件   2. 提取内容 3 .写入word

# 安装 python ——docx

from docx import Document
# 1.文档创建
document = Document（）
# 标题
document.add_heading("欢迎来到动脑学习python",level=1)
document.add_paragraph("人生苦短我用python")
para1 = document.add_paragraph("pycharm真强大")
para1.add_run('的确如此！')

# 插入图片
document.add_picture("kt.jpg",width=Inches(6))
# 插入表格
tb = document.add_table(rows=5, cols=3)

import random
f_row = tb.row[0].cells
f_row[0].text = "python"
f_row[1].text = "c"
f_row[2].text = "c++"
for i in range(1,5):
    


document.save("dragon.docx")